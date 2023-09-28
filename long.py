from google.colab import files
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Function to capitalize the first letter of each word in a sentence
def capitalize_first_letters(sentence):
    return ' '.join(word.capitalize() for word in sentence.split())

# Function to create a custom style for the table of contents
def create_toc_style(doc):
    style = doc.styles.add_style('TOCStyle', 1)  # Use '1' for paragraph style
    style.base_style = doc.styles['Normal']  # You can use 'Normal' as the base style
    style.font.size = Pt(12)
    style.font.bold = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.first_line_indent = Pt(0)
    style.paragraph_format.left_indent = Pt(0)
    style.paragraph_format.right_indent = Pt(0)
    return style

# Upload the Word document in Google Colab
uploaded_files = files.upload()

# Process the uploaded Word document
if uploaded_files:
    uploaded_file_name = list(uploaded_files.keys())[0]

    # Load the Word document
    doc = Document(uploaded_file_name)

    # Initialize a variable to keep track of the page number
    page_number = 1

    for paragraph in doc.paragraphs:
        # Split the paragraph into sentences
        sentences = paragraph.text.split('. ')

        # Check if there is at least one sentence in the paragraph
        if sentences:
            # Capitalize the first sentence and make it a header
            first_sentence = sentences[0]
            if first_sentence.strip():
                run = paragraph.runs[0]  # Access the first run in the paragraph
                run.font.size = Pt(16)  # Adjust font size as needed
                run.bold = True  # Make the text bold
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Print the formatted header with page number
                formatted_sentence = capitalize_first_letters(first_sentence)
                print(f'Page {page_number}: {formatted_sentence}\n')

                # Increment the page number
                page_number += 1

    # Create a table for the table of contents
    table_of_contents = doc.add_table(rows=1, cols=2)
    table_of_contents.allow_autofit = False
    table_of_contents.style = 'Table Grid'

    # Set column widths for the table
    table_of_contents.columns[0].width = Pt(50)
    table_of_contents.columns[1].width = Pt(300)

    # Add a title to the table of contents
    title_cell = table_of_contents.cell(0, 0)
    title_cell.text = 'Table of Contents'
    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_cell.paragraphs[0].bold = True

    # Create the table of contents entries
    for i in range(1, page_number):
        toc_entry = capitalize_first_letters(doc.paragraphs[i].text.split(". ")[0])
        toc_page_number = str(i)

        # Add the entry and page number to the table
        row = table_of_contents.add_row().cells
        row[0].text = toc_page_number
        row[1].text = toc_entry

    # Save the modified document to your local system
    output_file_name = 'modified_document_with_toc.docx'  # Replace with your desired file name
    doc.save(output_file_name)

    # Download the modified document to your local system
    files.download(output_file_name)